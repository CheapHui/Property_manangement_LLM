import os
import json
import requests
from typing import List
from pathlib import Path
from pypdf import PdfReader
from typing import Optional
import subprocess
from docx import Document
import re
from dotenv import load_dotenv

# Load environment variables from .env file
ROOT = Path(__file__).resolve().parent.parent
load_dotenv(ROOT.parent / ".env")

# ========= API Configuration (loaded from .env) =========
DEEPSEEK_API_KEY = os.getenv("OPENAI_API_KEY")  # DeepSeek API key from .env
BASE_URL = os.getenv("OPENAI_API_BASE", "https://api.deepseek.com")
MODEL_NAME = "deepseek-chat"  # or "deepseek-reasoner"
# =================================
CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]")

def antiword_extract(doc_path: str, timeout_sec: int = 30) -> str:
    doc_abs = str(Path(doc_path).resolve())
    cmd = ["antiword", doc_abs]
    proc = subprocess.run(cmd, check=True, capture_output=True, text=True, timeout=timeout_sec)
    return proc.stdout or ""

def lo_cat_text(doc_path: str, timeout_sec: int = 90) -> str:
    soffice = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    doc_abs = str(Path(doc_path).resolve())

    cmd = [
        soffice,
        "--headless",
        "--nologo",
        "--nodefault",
        "--norestore",
        "--nolockcheck",
        "--cat",   # ✅ 直接 dump text
        doc_abs
    ]

    print("🛠️ LO --cat:", Path(doc_abs).name)
    proc = subprocess.run(
        cmd,
        check=True,
        capture_output=True,
        text=True,
        timeout=timeout_sec
    )
    return proc.stdout or ""

def docx_to_text(docx_path: str) -> str:
    doc = Document(docx_path)
    parts = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                ct = (cell.text or "").strip().replace("\n", " ")
                if ct:
                    row_text.append(ct)
            if row_text:
                parts.append(" | ".join(row_text))

    return "\n\n".join(parts)


def doc_to_txt(
    doc_path: str,
    txt_path: str,
    encoding: str = "utf-8",
    temp_convert_dir: str = "./_tmp_doc_convert",
    timeout_sec: int = 60   # ✅ 加呢行
) -> str:
    """
    主路線：
    - .doc  : LibreOffice -> .txt（有 timeout，唔會卡死）
    - .docx : python-docx 抽 text
    """
    doc_path_p = Path(doc_path)
    ext = doc_path_p.suffix.lower()

    if ext == ".doc":
        out_txt = convert_doc_to_txt_via_lo(
            str(doc_path_p),
            temp_convert_dir,
            timeout_sec=timeout_sec   # ✅ 傳落去
        )
        with open(out_txt, "r", encoding=encoding, errors="ignore") as f:
            full_text = f.read()

    elif ext == ".docx":
        full_text = docx_to_text(str(doc_path_p))

    else:
        raise ValueError(f"Unsupported file type: {ext}")

    txt_path_p = Path(txt_path)
    txt_path_p.parent.mkdir(parents=True, exist_ok=True)
    with open(txt_path_p, "w", encoding=encoding) as f:
        f.write(full_text)

    print(f"✓ Saved TXT: {txt_path_p}")
    return str(txt_path_p)

def convert_doc_to_txt_via_lo(doc_path: str, out_dir: str, timeout_sec: int = 60) -> str:
    """
    主路線：LibreOffice convert-to txt (UTF-8)
    fallback：
      1) --cat dump text（較穩）
      2) antiword（第三層）
    """
    os.makedirs(out_dir, exist_ok=True)

    soffice = "/Applications/LibreOffice.app/Contents/MacOS/soffice"

    profile_dir = Path(out_dir) / "_lo_profile"
    profile_dir.mkdir(parents=True, exist_ok=True)
    user_install = profile_dir.resolve().as_uri()

    doc_abs = str(Path(doc_path).resolve())
    out_abs = str(Path(out_dir).resolve())

    cmd = [
        soffice,
        "--headless",
        "--nologo",
        "--nodefault",
        "--norestore",
        "--nolockcheck",
        f"-env:UserInstallation={user_install}",
        "--convert-to", "txt:Text (encoded):UTF8",
        "--outdir", out_abs,
        doc_abs,
    ]

    base = Path(doc_abs).stem
    out_txt = Path(out_abs) / f"{base}.txt"

    print("🛠️ LO convert:", Path(doc_abs).name, "-> txt")

    def _return_if_exists() -> Optional[str]:
        if out_txt.exists() and out_txt.stat().st_size > 0:
            return str(out_txt)
        candidates = list(Path(out_abs).glob(f"{base}*.txt"))
        if candidates and candidates[0].stat().st_size > 0:
            return str(candidates[0])
        return None

    try:
        subprocess.run(cmd, check=True, capture_output=True, text=True, timeout=timeout_sec)
        got = _return_if_exists()
        if got:
            return got
        raise FileNotFoundError(f"Converted txt not found for {doc_abs}")

    except (subprocess.TimeoutExpired, subprocess.CalledProcessError) as e:
        # 先試 --cat（通常救到一大批）
        if isinstance(e, subprocess.TimeoutExpired):
            print(f"⏳ convert-to timeout ({timeout_sec}s). Fallback to --cat: {Path(doc_abs).name}")
        else:
            print(f"⚠️ convert-to failed (exit={e.returncode}). Fallback to --cat: {Path(doc_abs).name}")

        try:
            text = lo_cat_text(doc_abs, timeout_sec=120).strip()
            if text:
                with open(out_txt, "w", encoding="utf-8") as f:
                    f.write(text)
                return str(out_txt)
        except Exception:
            pass

        # 再試 antiword（第三層）
        print(f"🧯 Fallback to antiword: {Path(doc_abs).name}")
        try:
            text = antiword_extract(doc_abs, timeout_sec=30).strip()
            if text:
                with open(out_txt, "w", encoding="utf-8") as f:
                    f.write(text)
                return str(out_txt)
        except Exception:
            pass

        # 最後寫 error log（如果係 CalledProcessError 才有 stdout/stderr）
        log_path = Path(out_abs) / f"{base}.lo_error.txt"
        with open(log_path, "w", encoding="utf-8") as f:
            f.write("CMD:\n" + " ".join(cmd) + "\n\n")
            if isinstance(e, subprocess.CalledProcessError):
                f.write("STDOUT:\n" + ((e.stdout or "").strip()) + "\n\n")
                f.write("STDERR:\n" + ((e.stderr or "").strip()) + "\n")
            else:
                f.write(f"ERROR:\nTimeout after {timeout_sec}s\n")

        raise RuntimeError(f"LibreOffice convert failed. See: {log_path}") from e
    
def is_case_done(jsonl_path: Path, min_lines: int = 1) -> bool:
    """
    判斷 case 是否已完成：
    - 檔案存在
    - 檔案 size > 0
    - 至少有 min_lines 行（每行一個 chunk）
    """
    if not jsonl_path.exists():
        return False
    if jsonl_path.stat().st_size <= 0:
        return False
    try:
        with jsonl_path.open("r", encoding="utf-8") as f:
            cnt = 0
            for _ in f:
                cnt += 1
                if cnt >= min_lines:
                    return True
        return False
    except Exception:
        return False


def mark_case_failed(jsonl_path: Path, reason: str):
    """
    失敗時將輸出檔改名做 .failed，避免下次誤判為 done。
    """
    try:
        failed_path = jsonl_path.with_suffix(jsonl_path.suffix + ".failed")
        if jsonl_path.exists():
            jsonl_path.rename(failed_path)
        # 同時寫原因
        with open(str(failed_path) + ".reason.txt", "w", encoding="utf-8") as f:
            f.write(reason)
    except Exception:
        pass

def sanitize_for_json(s: str) -> str:
    """
    移除 JSON string 不允許的控制字元
    (保留 \n \r \t)
    """
    if not s:
        return s
    # 先統一換行
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    # 移除其餘控制字元
    s = CONTROL_CHARS_RE.sub("", s)
    return s

def docx_to_text(docx_path: str) -> str:
    doc = Document(docx_path)
    parts = []

    # paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # tables (table 文字都好重要)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                ct = (cell.text or "").strip()
                ct = ct.replace("\n", " ").strip()
                row_text.append(ct)
            line = " | ".join([x for x in row_text if x])
            if line:
                parts.append(line)

    return "\n\n".join(parts)


def convert_doc_to_docx(doc_path: str, out_dir: str) -> str:
    """
    用 LibreOffice 將 .doc 轉 .docx
    - 需要已安裝 LibreOffice，並且 soffice 在 PATH 內（或你改成完整路徑）
    """
    os.makedirs(out_dir, exist_ok=True)

    # Windows 例子：如果你 PATH 無 soffice，可改成：
    # soffice = r"C:\Program Files\LibreOffice\program\soffice.exe"
    soffice = "soffice"

    cmd = [
        soffice,
        "--headless",
        "--convert-to", "docx",
        "--outdir", out_dir,
        doc_path
    ]
    subprocess.run(cmd, check=True,timeout=TimeoutError)

    base = Path(doc_path).stem
    converted = str(Path(out_dir) / f"{base}.docx")
    if not os.path.exists(converted):
        raise FileNotFoundError(f"LibreOffice convert failed, docx not found: {converted}")
    return converted



def pdf_to_txt(
    pdf_path: str,
    txt_path: Optional[str] = None,
    skip_first_pages: int = 0,
    max_pages: Optional[int] = None,
    encoding: str = "utf-8"
) -> str:
    """
    將單一 PDF 轉做 TXT。

    :param pdf_path: PDF 檔路徑
    :param txt_path: 輸出 TXT 檔路徑，如為 None 則用同名 .txt
    :param skip_first_pages: 跳過頭幾頁（例如 index / 封面）
    :param max_pages: 最多處理幾多頁（None = 全部）
    :param encoding: 輸出 TXT 編碼
    :return: 輸出 TXT 檔實際路徑
    """
    pdf_path = Path(pdf_path)

    reader = PdfReader(str(pdf_path))

    num_pages = len(reader.pages)
    start_page = min(skip_first_pages, num_pages)
    if max_pages is not None:
        end_page = min(start_page + max_pages, num_pages)
    else:
        end_page = num_pages

    print(f"Reading PDF: {pdf_path}")
    print(f"Total pages: {num_pages}, processing pages: {start_page} to {end_page - 1}")

    texts = []
    for i in range(start_page, end_page):
        page = reader.pages[i]
        try:
            page_text = page.extract_text() or ""
        except Exception as e:
            print(f"⚠️ Page {i} extract_text() error: {e}")
            page_text = ""
        texts.append(page_text)

    full_text = "\n\n==== PAGE BREAK ====\n\n".join(texts)
    
    # Save to file if txt_path is provided
    if txt_path:
        txt_path = Path(txt_path)
        txt_path.parent.mkdir(parents=True, exist_ok=True)
        with open(txt_path, "w", encoding=encoding) as f:
            f.write(full_text)
        print(f"✓ Saved TXT: {txt_path}")
        return str(txt_path)
    
    return full_text

SYSTEM_PROMPT = """
你是一名精通香港物業管理及法律文件處理的法律助理，負責把香港法院判詞轉換成適合 RAG 使用的 JSON 結構。

輸出要求：
- 只輸出 JSON（不包含任何解釋、前文後理、標題或註解）
- JSON 的 key 名必須與我提供的 schema 完全一致
- 如某些欄位無法從文本判斷，請設為 null 或空陣列 []，不要亂猜案件內容
- 用香港中文及繁體字回答
- 你必需只可以選用我提供的 property_topic, ordinance_sections, building_type, roles_involved, outcome 選項
- 輸出格式：
- MUST 輸出一個 JSON object
- JSON object 格式必須係：
  {
    "chunks": [ ...chunk objects... ]
  }
- 不要輸出其他任何文字
- 輸出 JSON 時，請移除所有不可見控制字元（除 \n, \t 外）。
- 每次最多輸出 4 個 chunk objects。若文本太長，只輸出最重要的 4 個。

----------------------------------------------------------------
### METADATA VOCABULARY (STRICT)
----------------------------------------------------------------

## property_topic (use ONLY values below; may choose 1–3 items)
OC_powers
OC_litigation_rights
OC_meeting_procedures
OC_resolution_validity
OC_election_and_officers
management_fee
special_levy
sinking_fund
budget_and_accounts
building_maintenance
common_parts_dispute
water_seepage
unauthorized_structure
fire_safety
nuisance
use_restrictions
access_dispute
parking_and_vehicles
DMC_interpretation
manager_role_under_DMC
management_company_liability
service_contracts
BMO_interpretation
land_lease_and_title
government_enforcement
land_tribunal_procedure

## roles_involved (choose all roles that appear)
Owners_Corporation
Individual_Owner
Joint_Owners
Management_Company
Manager_under_DMC
Tenant
Developer
Contractor
Government_Department
Incorporated_Owners_Other_Building

## ordinance_sections (ONLY values matching this format):
<ShortName>_Cap<Number>_s<Section>
Examples:
BMO_Cap344_s16
BMO_Cap344_s18
BMO_Cap344_s34I
BO_Cap123_s24
FSO_Cap95_s9
LTO_Cap219_s41

## building_type (choose 0–1)
residential
composite
industrial
commercial
shopping_centre
village_house
mixed_use_estate

## outcome (choose 1)
Plaintiff_succeeds
Defendant_succeeds
Mixed
Dismissed
Settlement
Appeal_allowed
Appeal_dismissed
null

## section_type (choose 1 per chunk)
headnote
facts
issues
reasoning
order
procedural_history

你要處理的是「單一宗 court case」，輸入會包含：
- 基本資料（court, case no., parties, date 等，如有）
- 判詞內文（可能只係其中一部分）

你需要：
1. 分析案件大意（與香港物業管理相關的重點，例如：管理費、維修費、DMC 解釋、OC 權力、管理公司責任等）
2. 把判詞切成多個 content chunk，每個 chunk 代表一段重要 reasoning / 解釋文本
3. 為每個 chunk 輸出一行 JSON（list 形式），用以下 schema：

每一個 chunk object 應該有以下欄位：

{
  "id": "case_<CASE_ID>_chunk1",
  "case_id": "HKCFI_2020_123",
  "source": "HKCFI",
  "case_no": "HCMP 123/2020",
  "court": "Court of First Instance",
  "judgment_date": "2020-05-12",
  "year": 2020,
  "parties": "Owners' Corporation of XXX Building v YYY Management Co Ltd",
  "url": null,

  "property_topic": [], // Only from vocabulary above
  "ordinance_sections": [], // ONLY from vocab or empty
  "building_type": [], // ONLY from vocab or empty
  "roles_involved": [], ONLY from vocab
  "outcome": null, //MUST use vocab
  "is_leading_case": false,

  "paragraph_range": "[12]-[18]",
  "content": "這個 chunk 的原文內容...",

  "content_summary": "用 2-4 句中文總結這段在物業管理上的重點。",
  "issue_summary": "用 1-2 句中文概括這段涉及的法律爭點。",

  "category": "court_case",
  "section_type": "reasoning",

  "embedding_model": null
}

輸出格式：
- MUST 輸出一個 JSON array（list），每個元素係一個 chunk object
- 不要輸出其他任何文字
"""

# ===============================
# 1. 將好長嘅判詞 TEXT 先切成幾份 sub-text
# ===============================

def split_text_into_chunks(text: str,
                           max_chars: int = 8000,
                           overlap_chars: int = 500) -> List[str]:
    """
    用段落 (\n\n) 為單位，將長文本切成多個 chunk。
    - max_chars: 每個 chunk 目標最大字元數（粗略近似 token limit）
    - overlap_chars: chunk 與 chunk 之間保留多少字元 overlap（用最後幾個段落）
    """
    paragraphs = text.split("\n\n")
    chunks = []
    current = []
    current_len = 0

    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        p_len = len(p) + 2  # 加番 \n\n

        if current_len + p_len > max_chars and current:
            # 完成一個 chunk
            chunk_text = "\n\n".join(current)
            chunks.append(chunk_text)

            # 準備下一個 chunk：加 overlap
            if overlap_chars > 0:
                # 由尾開始揀返啲 paragraph 做 overlap
                overlap = []
                overlap_len = 0
                for para in reversed(current):
                    l = len(para) + 2
                    if overlap_len + l > overlap_chars:
                        break
                    overlap.insert(0, para)
                    overlap_len += l
                current = overlap + [p]
                current_len = sum(len(x) + 2 for x in current)
            else:
                current = [p]
                current_len = p_len
        else:
            current.append(p)
            current_len += p_len

    if current:
        chunks.append("\n\n".join(current))

    return chunks


# ===============================
# 2. DeepSeek API：處理一個 sub-text → JSON chunks
# ===============================

def call_deepseek_court_case_to_chunks(case_text: str, retries: int = 2) -> List[dict]:
    """
    呼叫 DeepSeek，把一段 court case 文本轉做 chunk JSON list。
    - 有 retry：當模型輸出唔係合法 JSON array 時會重試
    """
    url = f"{BASE_URL}/v1/chat/completions"

    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }

    case_text = sanitize_for_json(case_text)

    payload = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": case_text}
        ],
        "stream": False,
        "temperature": 0.1,
        "max_tokens": 8192,
        "response_format": {"type": "json_object"}
    }

    last_err = None

    

    for attempt in range(retries + 1):
        try:
            resp = requests.post(url, headers=headers, json=payload, timeout=120)
            resp.raise_for_status()
            data = resp.json()

            content = data["choices"][0]["message"]["content"].strip()

            # 只擷取 JSON array（由第一個 [ 到最後一個 ]）
            start = content.find("[")
            end = content.rfind("]")
            if start != -1 and end != -1 and end > start:
                content_json = content[start:end + 1]
            else:
                content_json = content  # fallback

            content_json = sanitize_for_json(content_json)
            chunks = json.loads(content_json)

            if not isinstance(chunks, list):
                raise ValueError("Model output is not a JSON array (list).")
            

            return chunks



        except Exception as e:
            last_err = e
            # dump raw output（如果有）
            try:
                with open(f"deepseek_bad_output_attempt{attempt}.txt", "w", encoding="utf-8") as f:
                    f.write(content if "content" in locals() else str(e))
            except Exception:
                pass

            if attempt == retries:
                raise

            print(f"⚠️ DeepSeek output parse failed, retrying... attempt {attempt+1}/{retries}")

    # 理論上唔會到呢度
    raise last_err


# ===============================
# 3. 將「好多 sub-text」全部丟去 DeepSeek 再合併
# ===============================

def process_long_case_text_to_chunks(case_text: str,
                                     case_id: str,
                                     max_chars: int = 5000,
                                     overlap_chars: int = 200) -> List[dict]:
    """
    - 先將超長判詞切成多個 sub-text chunk
    - 逐個 sub-text 丟去 DeepSeek 做 JSON chunks
    - 最後合併晒，並統一重排 id：<case_id>_chunk1/2/3/...
    """
    # 1) 切 text
    text_chunks = split_text_into_chunks(
        case_text,
        max_chars=max_chars,
        overlap_chars=overlap_chars
    )

    print(f"分拆為 {len(text_chunks)} 個 sub-text chunk")

    all_chunks: List[dict] = []

    for idx, sub_text in enumerate(text_chunks, start=1):
        print(f"--> 處理 sub-text {idx}/{len(text_chunks)}")
        sub_chunks = call_deepseek_court_case_to_chunks(sub_text)
        all_chunks.extend(sub_chunks)

    # 2) 統一重設 id，同時確保 case_id 一致
    for i, ch in enumerate(all_chunks, start=1):
        # 如果 model 冇用你提供嘅 case_id，就以你傳入嘅為準
        ch["case_id"] = case_id
        ch["id"] = f"{case_id}_chunk{i}"

    return all_chunks


# ===============================
# 4. 寫入 JSONL 檔
# ===============================

def write_chunks_to_jsonl(chunks: List[dict], output_path: str):
    """
    將 chunks list 寫入 JSONL 檔，每行一個 JSON object。
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        for row in chunks:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")


# ===============================
# 5. main：示範處理單一宗 case
# ===============================

def main():
    input_pdf_folder = "./Count_case"
    intermediate_txt_folder = "../data/raw/cases_txt"
    output_jsonl_folder = "./data/chunked/court_cases"

    pdf_folder = Path(input_pdf_folder)
    txt_folder = Path(intermediate_txt_folder)
    jsonl_folder = Path(output_jsonl_folder)

    txt_folder.mkdir(parents=True, exist_ok=True)
    jsonl_folder.mkdir(parents=True, exist_ok=True)

    # 找所有 DOC / DOCX
    doc_files = sorted(
        f for f in pdf_folder.iterdir()
        if f.is_file()
        and f.suffix.lower() in [".doc", ".docx"]
        and not f.name.startswith("~$")
    )
    print(f"📁 在資料夾找到 {len(doc_files)} 個 DOC/DOCX 檔案")

    for doc_file in doc_files:
        case_id = doc_file.stem
        txt_output_path = txt_folder / f"{case_id}.txt"
        jsonl_output_path = jsonl_folder / f"{case_id}.jsonl"

        # 已完成就 skip
        if is_case_done(jsonl_output_path, min_lines=1):
            print(f"⏭️ Skip (already ingested): {case_id}")
            continue

        # 如之前失敗過：你可選擇 skip / retry
        failed_marker = jsonl_output_path.with_suffix(".jsonl.failed")
        if failed_marker.exists():
            print(f"⚠️ Found failed marker, re-trying: {case_id}")

        print("\n=============================================")
        print(f"📄 開始處理：{doc_file.name}")
        print("=============================================")

        try:
            # 1) DOC/DOCX → TXT（主路線：doc 轉 txt）
            print("→ Step 1: DOC/DOCX ➜ TXT")
            doc_to_txt(
                doc_path=str(doc_file),
                txt_path=str(txt_output_path),
                temp_convert_dir=str(txt_folder / "_tmp_doc_convert"),
                timeout_sec=60
            )

            # 2) 讀入 TXT
            print("→ Step 2: 讀取 TXT")
            with open(txt_output_path, "r", encoding="utf-8") as f:
                case_text = f.read()

            # 3) AI chunking
            print("→ Step 3: DeepSeek AI chunking")
            chunks = process_long_case_text_to_chunks(
                case_text=case_text,
                case_id=case_id,
                max_chars=8000,
                overlap_chars=500
            )

            # 4) 寫入 JSONL：先寫 tmp 再 replace，避免半成品
            print(f"→ Step 4: 寫入 JSONL：{jsonl_output_path}")
            tmp_path = jsonl_output_path.with_suffix(".jsonl.tmp")
            write_chunks_to_jsonl(chunks, str(tmp_path))
            os.replace(tmp_path, jsonl_output_path)

            print(f"✅ 完成 {case_id}（共 {len(chunks)} 個 chunks）")

        except Exception as e:
            msg = str(e)

            if "timed out" in msg and "--cat" in msg:
                reason = "LO_CAT_TIMEOUT"
            elif "Unterminated string" in msg:
                reason = "DEEPSEEK_JSON_TRUNCATED"
            elif "Invalid control character" in msg:
                reason = "DEEPSEEK_CONTROL_CHARS"
            else:
                reason = "OTHER"

            print(f"❌ Case failed: {case_id} — {reason} — {e}")
            mark_case_failed(jsonl_output_path, f"{reason}\n{e}")
            continue

    print("\n🎉 全部 DOC/DOCX 已處理完成！")

if __name__ == "__main__":
    main()
